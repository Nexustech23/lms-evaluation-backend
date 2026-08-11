# ============================================================
# DIAGRAM RENDERING ENGINE — Phase 3c
# Ported from controllers/institute/question_controller.py's diagram engine
# (matplotlib table/math/chemistry/graph/tree/network renderers + schemdraw
# circuits). One diagram DSL type — data_table — embeds as a native docx
# table (via embed_diagram); everything else renders to a PNG and is
# embedded as a picture with a caption.
#
# Security fix vs. Flask: the "graph" type's per-curve `expression` string
# (AI-generated, e.g. "sin(x)**2") was evaluated with Python's builtin
# eval() there (restricted globals/locals, but still arbitrary-AST exec).
# Here it goes through asteval.Interpreter, which restricts the AST itself
# (no import/exec/open/dunder-attribute access), not just the namespace.
# ============================================================

import logging
import math
import os
import tempfile
from io import BytesIO
from typing import Any, Dict

# matplotlib MUST be set to Agg before any other matplotlib/schemdraw import.
import matplotlib
matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import numpy as np
from asteval import Interpreter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    import schemdraw
    import schemdraw.elements as elm

    with schemdraw.Drawing(show=False) as _d:
        _d += elm.Line().right()
    _SCHEMDRAW_OK = True
    logging.info("schemdraw loaded OK")
except Exception as _sce:
    _SCHEMDRAW_OK = False
    logging.warning("schemdraw unavailable — circuits will use matplotlib fallback. Reason: %s", _sce)


def _fig_to_png(fig: "plt.Figure") -> bytes:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _schemdraw_to_png(drawing) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
        tmp = f.name
    try:
        drawing.save(tmp, dpi=150)
        with open(tmp, "rb") as f:
            return f.read()
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)


def _chem_unicode(text: str) -> str:
    import re

    sub_map = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
    sup_digits = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
    sup_sign = str.maketrans("+-", "⁺⁻")
    text = re.sub(r"<=>|<->", "⇌", text)
    text = re.sub(r"=>|->", "→", text)

    def _sup(m):
        digits = m.group(1) or ""
        sign = m.group(2) or ""
        return digits.translate(sup_digits) + sign.translate(sup_sign)

    text = re.sub(r"\^(\d*)([\+\-]?)", _sup, text)

    def _sub(m):
        return m.group(1) + m.group(2).translate(sub_map)

    text = re.sub(r"([A-Za-z\)])(\d+)", _sub, text)
    return text


def _render_data_table_png(spec: Dict[str, Any]) -> bytes:
    headers = spec.get("headers", [])
    rows = spec.get("rows", [])
    title = spec.get("title", "")
    n_cols = len(headers) or (len(rows[0]) if rows else 1)
    fig_h = max(2.0, 0.55 * (len(rows) + 1) + 0.8)
    fig, ax = plt.subplots(figsize=(max(7, n_cols * 1.6), fig_h))
    ax.axis("off")
    if title:
        fig.suptitle(title, fontsize=12, weight="bold", y=1.0)
    tbl = ax.table(cellText=rows, colLabels=headers, cellLoc="center", loc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.2, 1.7)
    for col in range(n_cols):
        cell = tbl[(0, col)]
        cell.set_facecolor("#1565c0")
        cell.set_text_props(color="white", weight="bold")
    for row in range(1, len(rows) + 1):
        bg = "#f0f4ff" if row % 2 == 0 else "#ffffff"
        for col in range(n_cols):
            tbl[(row, col)].set_facecolor(bg)
    return _fig_to_png(fig)


def _embed_docx_table(doc: Document, spec: Dict[str, Any]) -> None:
    headers = spec.get("headers", [])
    rows = spec.get("rows", [])
    title = spec.get("title", "")
    n_cols = len(headers) or (len(rows[0]) if rows else 1)

    if title:
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)

    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = hdr_row.cells[col_idx]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(str(header))
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1565c0")
        tc_pr.append(shd)

    for row_idx, row_data in enumerate(rows):
        fill = "f0f4ff" if row_idx % 2 == 0 else "ffffff"
        tbl_row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = tbl_row.cells[col_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

    doc.add_paragraph()


def _render_math_expression(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "")
    expressions = spec.get("expressions", [])
    if not expressions:
        return _render_biology_placeholder({"title": title or "Math Expression", "labels": {}})
    n = len(expressions)
    fig_h = max(1.5, 0.9 * n + 0.6)
    fig, ax = plt.subplots(figsize=(9, fig_h))
    ax.set_facecolor("#fafbff")
    fig.patch.set_facecolor("#fafbff")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=12, weight="bold", pad=8)
    step = 1.0 / (n + 1)
    for i, expr in enumerate(expressions):
        latex = expr.strip()
        if not (latex.startswith("$") and latex.endswith("$")):
            latex = f"${latex}$"
        y = 1.0 - (i + 1) * step
        try:
            ax.text(0.5, y, latex, ha="center", va="center", fontsize=15, transform=ax.transAxes)
        except Exception:
            ax.text(0.5, y, expr, ha="center", va="center", fontsize=13, transform=ax.transAxes)
    return _fig_to_png(fig)


def _render_chemical_equation(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Chemical Equations")
    equations = spec.get("equations", [])
    notes = spec.get("notes", [])
    if not equations:
        return _render_biology_placeholder({"title": title, "labels": {}})
    n = len(equations)
    fig_h = max(2.0, 0.85 * n + 0.9 + 0.4 * len(notes))
    fig, ax = plt.subplots(figsize=(9, fig_h))
    ax.set_facecolor("#f8fff8")
    fig.patch.set_facecolor("#f8fff8")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=13, weight="bold", pad=10)
    step = 1.0 / (n + len(notes) + 1)
    for i, eq in enumerate(equations):
        pretty = _chem_unicode(eq)
        y = 1.0 - (i + 1) * step
        ax.text(0.5, y, pretty, ha="center", va="center", fontsize=13, transform=ax.transAxes, fontfamily="DejaVu Sans")
        if i < n - 1:
            ax.axhline(y=y - step * 0.45, xmin=0.1, xmax=0.9, color="#ccddcc", linewidth=0.5)
    for j, note in enumerate(notes):
        y = 1.0 - (n + j + 1) * step
        ax.text(0.5, y, f"[{note}]", ha="center", va="center", fontsize=10, style="italic", color="#555555", transform=ax.transAxes)
    return _fig_to_png(fig)


_CIRCUIT_ELEMENT_MAP = {
    "resistor": lambda: elm.Resistor(),
    "capacitor": lambda: elm.Capacitor(),
    "inductor": lambda: elm.Inductor(),
    "battery": lambda: elm.BatteryCell(),
    "source": lambda: elm.SourceV(),
    "ground": lambda: elm.Ground(),
    "switch": lambda: elm.Switch(),
    "diode": lambda: elm.Diode(),
    "led": lambda: elm.LED(),
    "bulb": lambda: elm.Bulb(),
    "dot": lambda: elm.Dot(),
    "line": lambda: elm.Line(),
} if _SCHEMDRAW_OK else {}

_DIR_METHODS = {
    "right": lambda el: el.right(),
    "left": lambda el: el.left(),
    "up": lambda el: el.up(),
    "down": lambda el: el.down(),
}


def _render_electrical_circuit(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Circuit Diagram")
    elements = spec.get("elements", [])
    if not _SCHEMDRAW_OK or not elements:
        return _render_circuit_fallback(spec)
    try:
        with schemdraw.Drawing(show=False) as d:
            for el_spec in elements:
                etype = el_spec.get("type", "line").lower()
                if etype == "push":
                    d.push()
                    continue
                elif etype == "pop":
                    d.pop()
                    continue

                direction = el_spec.get("direction", "right").lower()
                label = el_spec.get("label", "")
                length = el_spec.get("length", None)
                factory = _CIRCUIT_ELEMENT_MAP.get(etype, _CIRCUIT_ELEMENT_MAP["line"])
                el = factory()
                dir_fn = _DIR_METHODS.get(direction, _DIR_METHODS["right"])
                el = dir_fn(el)
                if length is not None:
                    el = el.length(float(length))
                if label:
                    loc = el_spec.get("loc")
                    if not loc:
                        loc_map = {"up": "left", "down": "right", "right": "top", "left": "bottom"}
                        loc = loc_map.get(direction, "top")
                    ofst = el_spec.get("ofst", None)
                    if ofst is not None:
                        el = el.label(label, loc=loc, ofst=float(ofst))
                    else:
                        el = el.label(label, loc=loc)
                if el_spec.get("reverse"):
                    el = el.reverse()
                d += el

        png = _schemdraw_to_png(d)
        if title:
            img_arr = plt.imread(BytesIO(png))
            fig, ax = plt.subplots(figsize=(8, 5))
            ax.imshow(img_arr)
            ax.axis("off")
            ax.set_title(title, fontsize=13, weight="bold")
            return _fig_to_png(fig)
        return png
    except Exception as e:
        logging.error("schemdraw render failed: %s", e)
        return _render_circuit_fallback(spec)


def _render_circuit_fallback(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Circuit Diagram")
    elements = spec.get("elements", [])
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.set_facecolor("#f8faff")
    ax.axis("off")
    ax.set_title(title, fontsize=13, weight="bold")
    ax.plot([1, 9], [4.5, 4.5], "k-", lw=2)
    ax.plot([1, 9], [1.5, 1.5], "k-", lw=2)
    ax.plot([1, 1], [1.5, 4.5], "k-", lw=2)
    ax.plot([9, 9], [1.5, 4.5], "k-", lw=2)

    if not elements:
        ax.text(5.0, 3.0, "Circuit diagram elements unspecified.\nPlease refer to the values in the question text.",
                 ha="center", va="center", fontsize=10, color="#b71c1c", style="italic",
                 bbox=dict(boxstyle="round,pad=0.3", facecolor="#ffeacc", edgecolor="#ff9800", alpha=0.9))
    else:
        n = len(elements)
        spacing = 8.0 / n
        colors = ["#1565c0", "#2e7d32", "#c62828", "#6a1b9a", "#e65100"]
        for i, el in enumerate(elements):
            x = 1 + spacing * i + spacing / 2
            etype = el.get("type", "resistor").upper()
            el_label = el.get("label", "")
            display_text = f"{etype}\n({el_label})" if el_label else etype
            color = colors[i % len(colors)]
            rect = mpatches.FancyBboxPatch(
                (x - 0.55, 4.1), 1.1, 0.8,
                boxstyle="round,pad=0.05",
                linewidth=1.5, edgecolor=color, facecolor="white",
            )
            ax.add_patch(rect)
            ax.text(x, 4.5, display_text, ha="center", va="center", fontsize=8, weight="bold", color=color)

    ax.plot([0.7, 1.3], [3.0, 3.0], "k-", lw=4)
    ax.plot([0.85, 1.15], [2.7, 2.7], "k-", lw=2)
    ax.text(0.3, 3.0, "V_in", ha="center", va="center", fontsize=12, color="red", weight="bold")
    return _fig_to_png(fig)


def _render_biology_placeholder(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Diagram")
    labels = spec.get("labels", {})
    notes = spec.get("notes", [])
    n_labels = len(labels)
    fig_h = max(3.5, 2.5 + n_labels * 0.35)
    fig, ax = plt.subplots(figsize=(9, fig_h))
    ax.set_facecolor("#fff9f0")
    fig.patch.set_facecolor("#fff9f0")
    ax.axis("off")
    ax.set_title(title, fontsize=14, weight="bold", pad=12)
    ax.add_patch(mpatches.FancyBboxPatch(
        (0.05, 0.70), 0.90, 0.18,
        transform=ax.transAxes,
        boxstyle="round,pad=0.02",
        linewidth=1.5, edgecolor="#e65100", facecolor="#fff3e0",
    ))
    ax.text(0.50, 0.79,
            "⚠  Insert diagram manually during document editing",
            ha="center", va="center",
            fontsize=10, color="#bf360c", style="italic",
            transform=ax.transAxes)
    if labels:
        ax.text(0.05, 0.65, "Labels to include:", weight="bold", fontsize=10,
                transform=ax.transAxes, color="#333333")
        for i, (key, val) in enumerate(labels.items()):
            y = 0.58 - i * 0.09
            ax.text(0.08, y, f"  {key}.", fontsize=9, weight="bold", color="#1565c0", transform=ax.transAxes)
            ax.text(0.18, y, val, fontsize=9, color="#333333", transform=ax.transAxes)
    if notes:
        y_start = 0.55 - n_labels * 0.09
        ax.text(0.05, y_start, "Notes:", weight="bold", fontsize=10, transform=ax.transAxes, color="#333333")
        for j, note in enumerate(notes):
            ax.text(0.08, y_start - 0.09 * (j + 1), f"• {note}", fontsize=9, color="#555555", transform=ax.transAxes)
    return _fig_to_png(fig)


_CURVE_SYMBOLS = {
    "sin": np.sin, "cos": np.cos, "tan": np.tan, "sqrt": np.sqrt,
    "log": np.log, "exp": np.exp, "abs": np.abs, "pi": np.pi, "e": np.e,
    "asin": np.arcsin, "acos": np.arccos, "atan": np.arctan,
    "sinh": np.sinh, "cosh": np.cosh, "tanh": np.tanh,
}


def _eval_curve(expr: str, x: np.ndarray) -> np.ndarray:
    """AST-restricted curve evaluator (asteval) — replaces Flask's raw eval()."""
    aeval = Interpreter()
    aeval.symtable.update(_CURVE_SYMBOLS)
    aeval.symtable["x"] = x
    aeval.symtable["np"] = np
    result = aeval(expr)
    if aeval.error:
        raise ValueError(aeval.error[0].get_error()[1])
    return result


def _render_graph(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Graph")
    xlabel = spec.get("xlabel", "x")
    ylabel = spec.get("ylabel", "y")
    xrange = spec.get("xrange", [-10, 10])
    curves = spec.get("curves", [])
    x = np.linspace(float(xrange[0]), float(xrange[1]), 500)
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_facecolor("#f8f9fa")
    ax.grid(True, linestyle="--", alpha=0.5)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_title(title, fontsize=13, weight="bold")
    ax.set_xlabel(xlabel, fontsize=11)
    ax.set_ylabel(ylabel, fontsize=11)
    colors = ["#1565c0", "#c62828", "#2e7d32", "#6a1b9a", "#e65100"]

    if not curves:
        ax.text(0.5, 0.5, "No curves specified", ha="center", va="center", transform=ax.transAxes, fontsize=12, color="gray")
    else:
        plotted = False
        for i, curve in enumerate(curves):
            expr = curve.get("expression", "x").strip().replace("^", "**")
            label = curve.get("label", f"f{i+1}(x)")
            color = curve.get("color", colors[i % len(colors)])
            try:
                y = _eval_curve(expr, x)
                ax.plot(x, y, label=label, color=color, linewidth=2)
                plotted = True
            except Exception as e:
                logging.warning("Graph curve '%s' failed: %s", expr, e)
        if plotted:
            ax.legend(fontsize=10)
    return _fig_to_png(fig)


def _render_binary_tree(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Binary Tree")
    nodes_list = spec.get("nodes", [])

    nodes_map = {str(n["id"]): n for n in nodes_list}

    children = set()
    for n in nodes_list:
        if n.get("left"):
            children.add(str(n["left"]))
        if n.get("right"):
            children.add(str(n["right"]))

    all_ids = set(nodes_map.keys())
    roots = list(all_ids - children)
    root = roots[0] if roots else (nodes_list[0]["id"] if nodes_list else None)

    if not root:
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.axis("off")
        ax.text(0.5, 0.5, "Empty Binary Tree", ha="center", va="center")
        return _fig_to_png(fig)

    pos: Dict[str, tuple] = {}

    def get_height(node_id):
        if not node_id or node_id not in nodes_map:
            return 0
        node = nodes_map[node_id]
        return max(get_height(node.get("left")), get_height(node.get("right"))) + 1

    tree_height = get_height(root)

    def compute_positions(node_id, x, y, dx):
        if not node_id or node_id not in nodes_map:
            return
        pos[node_id] = (x, y)
        node = nodes_map[node_id]
        if node.get("left"):
            compute_positions(str(node["left"]), x - dx, y - 1, dx / 2)
        if node.get("right"):
            compute_positions(str(node["right"]), x + dx, y - 1, dx / 2)

    initial_dx = 2.0 ** (tree_height - 2) if tree_height > 1 else 1.0
    compute_positions(root, 0.0, 0.0, initial_dx)

    fig, ax = plt.subplots(figsize=(8, max(4, tree_height * 1.2)))
    ax.axis("off")
    ax.set_title(title, fontsize=12, weight="bold")

    for node_id, (x, y) in pos.items():
        node = nodes_map[node_id]
        if node.get("left") and str(node["left"]) in pos:
            cx, cy = pos[str(node["left"])]
            ax.plot([x, cx], [y, cy], "k-", zorder=1, lw=1.5)
        if node.get("right") and str(node["right"]) in pos:
            cx, cy = pos[str(node["right"])]
            ax.plot([x, cx], [y, cy], "k-", zorder=1, lw=1.5)

    for node_id, (x, y) in pos.items():
        circle = mpatches.Circle((x, y), 0.25, edgecolor="#1e40af", facecolor="#eff6ff", lw=2, zorder=2)
        ax.add_patch(circle)
        ax.text(x, y, node_id, ha="center", va="center", fontsize=10, weight="bold", color="#1e3a8a", zorder=3)

    xs = [x for x, y in pos.values()]
    ys = [y for x, y in pos.values()]
    ax.set_xlim(min(xs) - 0.5, max(xs) + 0.5)
    ax.set_ylim(min(ys) - 0.5, max(ys) + 0.5)

    return _fig_to_png(fig)


def _render_network_graph(spec: Dict[str, Any]) -> bytes:
    title = spec.get("title", "Graph")
    nodes = spec.get("nodes", [])
    edges = spec.get("edges", [])
    directed = spec.get("directed", False)

    if not nodes:
        node_set = set()
        for edge in edges:
            node_set.add(str(edge.get("from")))
            node_set.add(str(edge.get("to")))
        nodes = list(node_set)

    n_nodes = len(nodes)
    if n_nodes == 0:
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.axis("off")
        ax.text(0.5, 0.5, "Empty Graph", ha="center", va="center")
        return _fig_to_png(fig)

    pos = {}
    for i, node in enumerate(nodes):
        angle = 2 * math.pi * i / n_nodes
        pos[str(node)] = (math.cos(angle), math.sin(angle))

    fig, ax = plt.subplots(figsize=(6, 6))
    ax.axis("off")
    ax.set_title(title, fontsize=12, weight="bold")

    for edge in edges:
        u = str(edge.get("from"))
        v = str(edge.get("to"))
        weight = edge.get("weight")

        if u in pos and v in pos:
            x1, y1 = pos[u]
            x2, y2 = pos[v]
            dx, dy = x2 - x1, y2 - y1
            dist = math.sqrt(dx * dx + dy * dy)
            if dist > 0:
                sx, sy = x1 + (dx / dist) * 0.2, y1 + (dy / dist) * 0.2
                ex, ey = x2 - (dx / dist) * 0.2, y2 - (dy / dist) * 0.2
            else:
                sx, sy, ex, ey = x1, y1, x2, y2

            if directed:
                ax.annotate("", xy=(ex, ey), xytext=(sx, sy),
                            arrowprops=dict(arrowstyle="->", color="#475569", lw=1.5, shrinkA=0, shrinkB=0))
            else:
                ax.plot([sx, ex], [sy, ey], color="#475569", lw=1.5, zorder=1)

            if weight is not None:
                mx, my = (x1 + x2) / 2, (y1 + y2) / 2
                ox = -dy / dist * 0.1 if dist > 0 else 0
                oy = dx / dist * 0.1 if dist > 0 else 0
                ax.text(mx + ox, my + oy, str(weight), ha="center", va="center", fontsize=9, color="#b91c1c", weight="bold", zorder=2,
                         bbox=dict(boxstyle="round,pad=0.2", facecolor="#ffffff", edgecolor="none", alpha=0.8))

    for node in nodes:
        x, y = pos[str(node)]
        circle = mpatches.Circle((x, y), 0.2, edgecolor="#0f766e", facecolor="#f0fdf4", lw=2, zorder=3)
        ax.add_patch(circle)
        ax.text(x, y, str(node), ha="center", va="center", fontsize=10, weight="bold", color="#115e59", zorder=4)

    ax.set_xlim(-1.4, 1.4)
    ax.set_ylim(-1.4, 1.4)
    return _fig_to_png(fig)


def draw_diagram(spec: Dict[str, Any]) -> bytes:
    """Dispatch on spec['type'] -> PNG bytes. Never raises — falls back to an error image."""
    dtype = str(spec.get("type", "generic")).lower()
    try:
        if dtype == "data_table":
            return _render_data_table_png(spec)
        elif dtype == "math_expression":
            return _render_math_expression(spec)
        elif dtype in ("chemical_equation", "chemistry"):
            return _render_chemical_equation(spec)
        elif dtype in ("electrical_circuit", "circuit"):
            return _render_electrical_circuit(spec)
        elif dtype == "graph":
            return _render_graph(spec)
        elif dtype in ("binary_tree", "tree", "binary_search_tree"):
            return _render_binary_tree(spec)
        elif dtype in ("network_graph", "directed_graph", "undirected_graph"):
            return _render_network_graph(spec)
        else:
            return _render_biology_placeholder(spec)
    except Exception as e:
        logging.error("draw_diagram '%s' failed: %s", dtype, e)
        fig, ax = plt.subplots(figsize=(6, 2))
        ax.axis("off")
        ax.text(0.5, 0.5, f"Diagram error: {e}", ha="center", va="center",
                 fontsize=10, color="#b71c1c",
                 bbox=dict(boxstyle="round", facecolor="#fff3e0"))
        return _fig_to_png(fig)


_WIDTH_MAP = {
    "math_expression": Inches(5.5),
    "chemical_equation": Inches(6.0),
    "electrical_circuit": Inches(5.5),
    "graph": Inches(6.0),
}


def embed_diagram(doc: Document, spec: Dict[str, Any]) -> None:
    """Render a diagram spec into `doc` — native table for data_table, PNG + caption otherwise."""
    dtype = spec.get("type", "")
    if dtype == "data_table":
        _embed_docx_table(doc, spec)
        return
    try:
        png_bytes = draw_diagram(spec)
        buf = BytesIO(png_bytes)
        width = _WIDTH_MAP.get(dtype, Inches(5.0))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(buf, width=width)
        caption = spec.get("title", "") or dtype.replace("_", " ").title()
        if caption:
            cp = doc.add_paragraph(f"Figure: {caption}")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(8)
            for run in cp.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
    except Exception as e:
        logging.error("embed_diagram failed for type '%s': %s", dtype, e)
        p = doc.add_paragraph()
        r = p.add_run(f"[Diagram render error ({dtype}): {e}]")
        r.font.color.rgb = RGBColor(180, 0, 0)
        r.font.size = Pt(10)
