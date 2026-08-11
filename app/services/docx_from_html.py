# ============================================================
# Editor HTML -> styled DOCX (used by question-paper create/update).
# Ported from question_controller.py's _generate_docx_from_html.
# Blocking — run via asyncio.to_thread() from async callers.
# ============================================================

import base64
import logging
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, Optional

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.services.docx_common import (
    C_BLUE_DARK,
    C_BLUE_DEEP,
    C_BLUE_MID,
    C_TEXT,
    add_footer_with_page_number,
    build_exam_header,
    set_cell_bg,
    set_cell_border_all,
)


def generate_docx_from_html(html_content: str, header_data: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Converts editor HTML to a styled DOCX matching the preview pane:
    exam header, h1/h2/h3, p, ul/ol, table, blockquote, img (real embedding,
    not a stub), footer with page number.
    """
    if not html_content:
        raise ValueError("Editor content is empty")

    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)

    def _add_inline_content(para, node, base_size=11, base_color=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                txt = str(child)
                if txt:
                    r = para.add_run(txt)
                    r.font.size = Pt(base_size)
                    if base_color:
                        r.font.color.rgb = base_color
            else:
                r = para.add_run(child.get_text())
                r.font.size = Pt(base_size)
                if base_color:
                    r.font.color.rgb = base_color
                tag = child.name.lower() if child.name else ""
                if tag in ("strong", "b"):
                    r.bold = True
                if tag in ("em", "i"):
                    r.italic = True
                if tag == "u":
                    r.underline = True

    build_exam_header(doc, header_data)

    soup = BeautifulSoup(html_content, "html.parser")
    container = soup.body if soup.body else soup

    def render_image(el):
        src = el.get("src", "")
        if not src:
            return
        try:
            import requests
            from PIL import Image

            if src.startswith("data:image/") and ";base64," in src:
                _, base64_data = src.split(";base64,", 1)
                img_bytes = base64.b64decode(base64_data)
            elif src.startswith("http"):
                img_bytes = requests.get(src, timeout=10).content
            else:
                return

            img_pil = Image.open(BytesIO(img_bytes))
            out_io = BytesIO()
            img_pil.save(out_io, format="PNG")
            png_bytes = out_io.getvalue()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.add_run().add_picture(BytesIO(png_bytes), width=Inches(5.0))
        except Exception:
            logging.exception("Failed to embed image from html to docx")

    def render_block(el):
        tag = el.name.lower() if el.name else ""

        if tag == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(el.get_text().upper())
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = C_BLUE_DARK
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "4")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "bfdbfe")
            pBdr.append(bot)
            pPr.append(pBdr)

        elif tag == "h2":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(11)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(el.get_text().upper())
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = C_BLUE_MID

        elif tag == "h3":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(el.get_text())
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = C_BLUE_DEEP

        elif tag == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.line_spacing = Pt(20)
            _add_inline_content(p, el, base_size=11, base_color=C_TEXT)

        elif tag == "ul":
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = Pt(20)
                _add_inline_content(p, li, base_size=11, base_color=C_TEXT)

        elif tag == "ol":
            for i, li in enumerate(el.find_all("li", recursive=False), 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = Pt(20)
                r = p.add_run(f"{i}.  ")
                r.bold = True
                r.font.size = Pt(11)
                _add_inline_content(p, li, base_size=11, base_color=C_TEXT)

        elif tag == "table":
            rows_els = el.find_all("tr")
            if not rows_els:
                return
            n_cols = max(
                sum(int(td.get("colspan", 1)) for td in row.find_all(["td", "th"])) for row in rows_els
            )
            if n_cols == 0:
                return
            tbl = doc.add_table(rows=len(rows_els), cols=n_cols)
            tbl.style = "Table Grid"
            for ri, row_el in enumerate(rows_els):
                cells = row_el.find_all(["td", "th"])
                is_header = any(c.name == "th" for c in cells)
                for ci, cell_el in enumerate(cells[:n_cols]):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    if is_header:
                        set_cell_bg(cell, "eff6ff")
                        set_cell_border_all(cell, color="bfdbfe")
                        cr = cp.add_run(cell_el.get_text(strip=True))
                        cr.bold = True
                        cr.font.size = Pt(10)
                        cr.font.color.rgb = C_BLUE_DARK
                    else:
                        fill = "f8fafc" if ri % 2 == 0 else "ffffff"
                        set_cell_bg(cell, fill)
                        set_cell_border_all(cell, color="e2e8f0")
                        cr = cp.add_run(cell_el.get_text(strip=True))
                        cr.font.size = Pt(10)
                        cr.font.color.rgb = C_TEXT
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        elif tag == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(el.get_text())
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(71, 85, 105)

        elif tag == "img":
            render_image(el)

        else:
            txt = el.get_text(strip=True)
            if txt:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(txt)
                r.font.size = Pt(11)
                r.font.color.rgb = C_TEXT

    def traverse(node):
        if isinstance(node, NavigableString):
            txt = str(node).strip()
            if txt:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(txt)
                r.font.size = Pt(11)
                r.font.color.rgb = C_TEXT
            return

        tag = node.name.lower() if node.name else ""

        if tag in ("h1", "h2", "h3", "p", "table", "ul", "ol", "blockquote", "img"):
            render_block(node)
            if tag != "img":
                for nested_img in node.find_all("img"):
                    render_image(nested_img)
            return

        if tag in ("li", "tr", "td", "th", "strong", "b", "em", "i", "u", "span", "br", "head", "style", "script"):
            return

        for child in node.children:
            traverse(child)

    traverse(container)

    add_footer_with_page_number(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def process_and_upload_base64_images(html_content: str, faculty_id: str) -> str:
    """
    Rewrites inline base64 <img> tags to ImageKit URLs before docx generation.
    Blocking — run via asyncio.to_thread() from async callers.
    """
    if not html_content:
        return html_content

    from app.services.imagekit import upload_file_to_imagekit

    soup = BeautifulSoup(html_content, "html.parser")
    changed = False

    for img in soup.find_all("img"):
        src = img.get("src", "")
        if not (src.startswith("data:image/") and ";base64," in src):
            continue
        try:
            header, b64data = src.split(";base64,", 1)
            mime_type = header.replace("data:", "")
            ext = mime_type.split("/")[-1] if "/" in mime_type else "png"
            image_bytes = base64.b64decode(b64data)
            ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S%f")
            upload = upload_file_to_imagekit(
                image_bytes, f"editor-image-{ts}.{ext}", "/editor-images", ["editor-image", faculty_id]
            )
            img["src"] = upload["url"]
            changed = True
        except Exception:
            logging.exception("Failed to upload base64 editor image")

    return str(soup) if changed else html_content
