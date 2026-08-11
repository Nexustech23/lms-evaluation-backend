# ============================================================
# Shared DOCX building blocks used by both docx_from_html.py (editor HTML ->
# docx) and docx_from_text.py (Claude plain-text -> docx). Flask's original
# had two near-identical header builders (_generate_docx_from_html's inline
# header vs _build_exam_header_docx) — consolidated here into one.
# ============================================================

from typing import Any, Dict, Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

C_BLUE_DARK = RGBColor(29, 78, 216)
C_BLUE_MID = RGBColor(30, 64, 175)
C_BLUE_DEEP = RGBColor(30, 58, 138)
C_TEXT = RGBColor(17, 24, 39)
C_FOOTER = RGBColor(120, 120, 120)
C_LABEL = RGBColor(100, 116, 139)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def set_cell_border_all(cell, color: str = "bfdbfe", sz: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_bdr = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tc_bdr.append(b)
    tc_pr.append(tc_bdr)


def set_cell_padding(cell, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)


def set_table_width_100pct(tbl) -> None:
    tbl_elem = tbl._tbl
    tbl_pr = tbl_elem.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tbl_pr)
    for old in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def remove_table_borders(tbl) -> None:
    tbl_elem = tbl._tbl
    tbl_pr = tbl_elem.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tbl_pr.append(borders)


def add_hr(doc, color: str = "bfdbfe", sz: str = "4", style: str = "single"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), style)
    bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def build_exam_header(doc, header_data: Optional[Dict[str, Any]] = None) -> None:
    """Institute banner, exam type, academic year, info grid, max marks, closing rule."""
    header_data = header_data or {}
    h_institute = header_data.get("institute", "INSTITUTE NAME")
    h_dept = header_data.get("department", "")
    h_exam_type = header_data.get("examType", "EXAMINATION")
    h_subject = header_data.get("subjectName", "")
    h_semester = header_data.get("semester", "")
    h_acad_year = header_data.get("academicYear", "")
    h_duration = header_data.get("duration", "")
    h_max_marks = header_data.get("totalMarks", "")

    accent_tbl = doc.add_table(rows=1, cols=1)
    set_table_width_100pct(accent_tbl)
    remove_table_borders(accent_tbl)
    tr_pr = accent_tbl.rows[0]._tr.get_or_add_trPr()
    trh = OxmlElement("w:trHeight")
    trh.set(qn("w:val"), "60")
    trh.set(qn("w:hRule"), "exact")
    tr_pr.append(trh)
    accent_cell = accent_tbl.rows[0].cells[0]
    set_cell_bg(accent_cell, "1d4ed8")
    accent_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    crest_p = doc.add_paragraph()
    crest_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crest_p.paragraph_format.space_before = Pt(6)
    crest_p.paragraph_format.space_after = Pt(2)
    crest_r = crest_p.add_run("- * -")
    crest_r.font.size = Pt(9)
    crest_r.font.color.rgb = C_BLUE_MID

    uni_p = doc.add_paragraph()
    uni_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uni_p.paragraph_format.space_before = Pt(2)
    uni_p.paragraph_format.space_after = Pt(2)
    uni_r = uni_p.add_run(h_institute.upper())
    uni_r.bold = True
    uni_r.font.size = Pt(14)
    uni_r.font.color.rgb = C_BLUE_DARK

    dept_p = doc.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_p.paragraph_format.space_before = Pt(0)
    dept_p.paragraph_format.space_after = Pt(6)
    dept_r = dept_p.add_run(h_dept.upper() if h_dept else "")
    dept_r.bold = True
    dept_r.font.size = Pt(10)
    dept_r.font.color.rgb = C_BLUE_MID

    exam_p = doc.add_paragraph()
    exam_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_p.paragraph_format.space_before = Pt(8)
    exam_p.paragraph_format.space_after = Pt(4)
    exam_r = exam_p.add_run(h_exam_type.upper())
    exam_r.bold = True
    exam_r.underline = True
    exam_r.font.size = Pt(13)
    exam_r.font.color.rgb = C_BLUE_DARK

    ay_p = doc.add_paragraph()
    ay_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ay_p.paragraph_format.space_before = Pt(6)
    ay_p.paragraph_format.space_after = Pt(6)
    ay_p.add_run("Academic Year: ").font.size = Pt(10)
    ay_bold = ay_p.add_run(h_acad_year or "2024 - 2025")
    ay_bold.bold = True
    ay_bold.font.size = Pt(10)
    ay_bold.font.color.rgb = C_BLUE_DARK

    grid_cols = []
    if h_subject:
        grid_cols.append(("SUBJECT", h_subject))
    if h_exam_type:
        grid_cols.append(("EXAM TYPE", h_exam_type))
    if h_duration:
        grid_cols.append(("DURATION", h_duration))
    if h_semester:
        grid_cols.append(("SEMESTER", h_semester))
    if not grid_cols:
        grid_cols = [("SUBJECT", "—"), ("EXAM TYPE", h_exam_type or "—"), ("DURATION", "—")]

    n_cols = len(grid_cols)
    info_tbl = doc.add_table(rows=2, cols=n_cols)
    remove_table_borders(info_tbl)
    set_table_width_100pct(info_tbl)

    for ci, (lbl, val) in enumerate(grid_cols):
        for ri in range(2):
            set_cell_bg(info_tbl.rows[ri].cells[ci], "eff6ff")

        if ci > 0:
            for ri in range(2):
                div_cell = info_tbl.rows[ri].cells[ci]
                tc_pr = div_cell._tc.get_or_add_tcPr()
                tc_bdr = OxmlElement("w:tcBorders")
                left_b = OxmlElement("w:left")
                left_b.set(qn("w:val"), "single")
                left_b.set(qn("w:sz"), "4")
                left_b.set(qn("w:space"), "0")
                left_b.set(qn("w:color"), "bfdbfe")
                tc_bdr.append(left_b)
                tc_pr.append(tc_bdr)

        lbl_cell = info_tbl.rows[0].cells[ci]
        set_cell_padding(lbl_cell, top=60, bottom=0, left=80, right=80)
        lp = lbl_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(lbl)
        lr.font.size = Pt(7.5)
        lr.font.color.rgb = C_LABEL
        lr.bold = True

        val_cell = info_tbl.rows[1].cells[ci]
        set_cell_padding(val_cell, top=0, bottom=60, left=80, right=80)
        vp = val_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(val)
        vr.bold = True
        vr.font.size = Pt(10)
        vr.font.color.rgb = C_BLUE_DARK

    mm_p = doc.add_paragraph()
    mm_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mm_p.paragraph_format.space_before = Pt(4)
    mm_p.paragraph_format.space_after = Pt(2)
    mm_lbl = mm_p.add_run("Maximum Marks: ")
    mm_lbl.font.size = Pt(10)
    mm_val = mm_p.add_run(str(h_max_marks) if h_max_marks else "—")
    mm_val.bold = True
    mm_val.font.size = Pt(10)
    mm_val.font.color.rgb = C_BLUE_DARK

    dbl_p = doc.add_paragraph()
    dbl_p.paragraph_format.space_before = Pt(4)
    dbl_p.paragraph_format.space_after = Pt(8)
    pPr = dbl_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "double")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1d4ed8")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_footer_with_page_number(doc, text: str = "Generated by Gradelytics  |  All rights reserved.") -> None:
    add_hr(doc, color="999999", sz="4", style="single")
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after = Pt(2)
    fr = fp.add_run(text)
    fr.font.size = Pt(9)
    fr.font.color.rgb = C_FOOTER
    fr.italic = True

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
