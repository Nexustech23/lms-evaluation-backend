# ============================================================
# Claude's plain-text question-paper output -> styled DOCX.
# Ported from question_controller.py's _build_docx.
# Blocking — run via asyncio.to_thread() from async callers.
#
# Diagram blocks (<<<DIAGRAM>>>...<<<END_DIAGRAM>>>) are parsed out of the
# question text and rendered via app/services/diagram_render.py (Phase 3c) —
# a real PNG/table embed, not a placeholder paragraph.
# ============================================================

import json
import re
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

from app.services.diagram_render import embed_diagram
from app.services.docx_common import (
    C_BLUE_DARK,
    add_footer_with_page_number,
    build_exam_header,
    remove_table_borders,
    set_cell_bg,
    set_cell_padding,
    set_table_width_100pct,
)

_SECTION_RE = re.compile(r"^(SECTION\s+[A-Z]|PART\s+[IVX]+|PART\s+[A-Z]|UNIT\s+\d+)", re.IGNORECASE)
_QUESTION_RE = re.compile(r"^(Q\d+[\.\)])\s+(.*)")
_SUBPART_RE = re.compile(r"^\s*(\([a-z]\)|\([ivxIVX]+\))\s+(.*)")
_MARKS_RE = re.compile(r"(\[\d+\s*Marks?\])", re.IGNORECASE)
_CO_TAG_RE = re.compile(r"(\[CO\d+(?:\s*,\s*CO\d+)*\])", re.IGNORECASE)
_NOTE_RE = re.compile(r"^(attempt|answer|note|all questions|each carries|students must)", re.IGNORECASE)
C_GREEN_DARK = RGBColor(21, 128, 61)


def _strip_markdown(text: str) -> str:
    """Mirrors question_controller.py's _strip_markdown."""
    text = re.sub(r"\*{3}(.*?)\*{3}", r"\1", text)
    text = re.sub(r"\*{2}(.*?)\*{2}", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_{3}(.*?)_{3}", r"\1", text)
    text = re.sub(r"_{2}(.*?)_{2}", r"\1", text)
    text = re.sub(r"(?<!\w)_(.*?)_(?!\w)", r"\1", text)
    text = re.sub(r"(?<!\w)\*+(?!\w)", "", text)
    text = re.sub(r"^>\s*", "", text)
    return text.strip()


def _add_co_table(doc: Document, co_list: List[Dict[str, Any]]) -> None:
    """Course Outcomes reference table, shown once after the exam header."""
    if not co_list:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("COURSE OUTCOMES")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = C_BLUE_DARK

    tbl = doc.add_table(rows=1 + len(co_list), cols=2)
    set_table_width_100pct(tbl)
    remove_table_borders(tbl)

    hdr = tbl.rows[0].cells
    for ci, text in enumerate(("CO Code", "Description")):
        set_cell_bg(hdr[ci], "eff6ff")
        set_cell_padding(hdr[ci], top=60, bottom=60, left=100, right=100)
        hp = hdr[ci].paragraphs[0]
        hr = hp.add_run(text)
        hr.bold = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = C_BLUE_DARK

    for ri, co in enumerate(co_list, start=1):
        code_cell, desc_cell = tbl.rows[ri].cells
        set_cell_padding(code_cell, top=50, bottom=50, left=100, right=100)
        set_cell_padding(desc_cell, top=50, bottom=50, left=100, right=100)
        cr = code_cell.paragraphs[0].add_run(co.get("co_code", ""))
        cr.bold = True
        cr.font.size = Pt(9)
        dr = desc_cell.paragraphs[0].add_run(co.get("description", ""))
        dr.font.size = Pt(9)

    tail = doc.add_paragraph()
    tail.paragraph_format.space_after = Pt(8)


def build_docx(
    paper_text: str,
    institute_name: str,
    department_name: str,
    subject_name: str,
    exam_type: str,
    semester: str,
    academic_year: str,
    total_marks: Any,
    duration: str,
    co_list: Optional[List[Dict[str, Any]]] = None,
) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)

    build_exam_header(doc, {
        "institute": institute_name,
        "department": department_name,
        "examType": exam_type,
        "subjectName": subject_name,
        "semester": semester,
        "academicYear": academic_year,
        "duration": duration,
        "totalMarks": total_marks,
    })

    _add_co_table(doc, co_list or [])

    def add_para(text: str, *, bold=False, italic=False, size=11, color: Optional[RGBColor] = None, indent=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if indent is not None:
            p.paragraph_format.left_indent = Inches(indent)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    def add_section_heading(text: str):
        from app.services.docx_common import add_hr
        add_hr(doc)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = C_BLUE_DARK

    def add_question(prefix: str, rest: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r_prefix = p.add_run(f"{prefix} ")
        r_prefix.bold = True
        r_prefix.font.size = Pt(11)

        tag_matches = sorted(
            list(_MARKS_RE.finditer(rest)) + list(_CO_TAG_RE.finditer(rest)),
            key=lambda m: m.start(),
        )
        last_end = 0
        for match in tag_matches:
            if match.start() > last_end:
                r = p.add_run(rest[last_end:match.start()])
                r.font.size = Pt(11)
            r_tag = p.add_run(match.group(1))
            r_tag.bold = True
            r_tag.font.size = Pt(11)
            r_tag.font.color.rgb = C_GREEN_DARK if match.re is _CO_TAG_RE else C_BLUE_DARK
            last_end = match.end()
        if last_end < len(rest):
            r = p.add_run(rest[last_end:])
            r.font.size = Pt(11)

    def add_subpart(label: str, rest: str):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(4)
        r_label = p.add_run(f"{label} ")
        r_label.bold = True
        r_label.font.size = Pt(11)
        r_rest = p.add_run(rest)
        r_rest.font.size = Pt(11)

    lines = paper_text.splitlines()

    skip_until_content = True
    in_instructions = False
    in_diagram_block = False
    diagram_buffer: list[str] = []
    instr_num = 1

    for raw_line in lines:
        line = _strip_markdown(raw_line)

        if in_diagram_block:
            if line.strip() == "<<<END_DIAGRAM>>>":
                in_diagram_block = False
                try:
                    spec = json.loads("\n".join(diagram_buffer))
                    embed_diagram(doc, spec)
                except json.JSONDecodeError as e:
                    add_para(f"[Diagram JSON error: {e}]", color=RGBColor(180, 0, 0), size=10)
                diagram_buffer = []
            else:
                diagram_buffer.append(raw_line)
            continue

        if line.strip() == "<<<DIAGRAM>>>":
            in_diagram_block = True
            diagram_buffer = []
            continue

        if not line.strip():
            if skip_until_content:
                continue
            in_instructions = False
            continue

        if skip_until_content:
            if line.strip().upper().startswith("INSTRUCTIONS:") or _SECTION_RE.match(line.strip()):
                skip_until_content = False
            else:
                continue

        if line.strip().upper().startswith("INSTRUCTIONS:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("INSTRUCTIONS:")
            r.bold = True
            r.font.color.rgb = C_BLUE_DARK
            in_instructions = True
            instr_num = 1
            continue

        if _SECTION_RE.match(line.strip()):
            in_instructions = False
            add_section_heading(line.strip())
            continue

        if _NOTE_RE.match(line.strip()):
            add_para(line.strip(), italic=True, size=10, color=RGBColor(107, 114, 128))
            continue

        question_match = _QUESTION_RE.match(line.strip())
        if question_match:
            in_instructions = False
            add_question(question_match.group(1), question_match.group(2))
            continue

        subpart_match = _SUBPART_RE.match(line)
        if subpart_match:
            add_subpart(subpart_match.group(1), subpart_match.group(2))
            continue

        if in_instructions:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{instr_num}. {line.strip()}")
            r.italic = True
            r.font.size = Pt(10)
            instr_num += 1
            continue

        add_para(line.strip())

    add_footer_with_page_number(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
